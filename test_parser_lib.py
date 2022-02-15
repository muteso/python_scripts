import re
import docx

class PivotErr(Exception):
    pass

def deb_get_qstn(path, exp=r'\d+ *?\.'):
    file = docx.Document(path)
    buf = []
    for line in file.paragraphs:
        style = line.style
        line = line.text.lstrip()
        if re.match(exp, line):
            if len(buf) > 0:
                with open("qstn.txt", "w") as f:
                    f.write(str(style) + str(buf[0][0]) + buf[0][1] + '\n')
                    for i in buf[1:]:
                        f.write(i)
                buf = []
            text = line.rstrip().split('.')
            buf.append([int(text[0]), text[1]])
            if text[0] == "368":
                print("LOOK!")
        elif re.match(r".+$", line):
            if line != "":
                text = line.rstrip()
                buf.append(f"\t{text}\n")

def get_qstn(path, exp=r'\d+ *?\.'):
    '''
    Returns the list in form [№ question, text] (parses TEXT).\n
    exp= takes a regulat expression for separating № question and text
    from each other.
    '''
    file = docx.Document(path)
    qstn = []
    buf = []
    for line in file.paragraphs:
        line = line.text.lstrip()
        if re.match(exp, line):
            if len(buf) > 0:
                qstn.append(buf)
                buf = []
            text = line.rstrip().split('.')
            buf.append([int(text[0]), text[1]])
        elif re.match(r".+$", line):
            if line != "":
                text = line.rstrip()
                buf.append(f"\t{text}\n")
    return qstn

def get_answ(path, exp=r'[–-]', pivot="row"):
    '''
    Returns the list in form [№ question, answer] (parses TABLE).\n
    exp= takes a regulat expression for separating № question and answer from each other.\n
    pivot= takes a "col" or "row" as main pivot direction.
    '''
    file = docx.Document(path)
    answ = []
    for table in file.tables:
        if pivot == 'col':
            piv = table.columns
        elif pivot == 'row':
            piv = table.rows
        else:
            raise PivotErr
        for unit in piv:
            for cell in unit.cells:
                for line in cell.paragraphs:
                    line = line.text
                    if line != '':
                        t_answ = line.replace(' ', '')
                        t_answ = re.split(exp, t_answ)
                        answ.append([int(t_answ[0]), t_answ[1]])
    return answ